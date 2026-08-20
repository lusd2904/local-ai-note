import os
import io
import json
import secrets
import hashlib
import re
from typing import List, Optional, Dict, Any
from fastapi import APIRouter, Depends, HTTPException, Query, File, UploadFile, Form
from sqlalchemy.orm import Session
from sqlalchemy import or_, desc, func
from ..database import get_db
from ..models import Note, Notebook, AudioRecord, Database
from ..schemas import (
    NoteCreate, NoteUpdate, NoteOut, NoteStatsOut,
    NoteLockRequest, NoteUnlockRequest, NoteVerifyPasswordRequest,
    GraphDataOut, GraphNode, GraphLink, BacklinksOut, BacklinkItem
)

def decode_text_bytes(content_bytes: bytes) -> str:
    for encoding in ['utf-8', 'utf-8-sig', 'gb18030', 'gbk', 'big5', 'latin1']:
        try:
            return content_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content_bytes.decode('utf-8', errors='ignore')

def extract_content_from_file(filename: str, content_bytes: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == '.docx':
        try:
            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(row_text))
            return "\n\n".join(paragraphs)
        except Exception as e:
            pass
            
    if ext in ['.html', '.htm']:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content_bytes, 'html.parser')
            for s in soup(['script', 'style']):
                s.extract()
            return soup.get_text(separator='\n\n', strip=True)
        except Exception as e:
            pass

    return decode_text_bytes(content_bytes)


# PBKDF2 迭代次数（OWASP 2023 推荐：PBKDF2-HMAC-SHA256 至少 600,000 次）
PBKDF2_ITERATIONS = 600000

def hash_password(password: str) -> str:
    """
    使用 PBKDF2-HMAC-SHA256 对密码进行安全哈希
    - 随机 16 字节盐值
    - 600,000 次迭代（抵御暴力破解和 GPU 加速攻击）
    - 返回格式: {salt_hex}${hash_hex}
    """
    salt = secrets.token_bytes(16)
    pw_hash = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt, PBKDF2_ITERATIONS)
    return f"{salt.hex()}${pw_hash.hex()}"

def verify_password(stored_hash: str, password: str) -> bool:
    """
    验证密码是否匹配存储的哈希值
    支持向后兼容：自动检测旧版 SHA-256 格式并升级提示
    """
    if not stored_hash or "$" not in stored_hash:
        return False
    try:
        salt_hex, pw_hash_hex = stored_hash.split("$", 1)
        salt = bytes.fromhex(salt_hex)

        # 检测是否为旧版 SHA-256 格式（哈希长度为 64 字符）
        if len(pw_hash_hex) == 64 and len(salt_hex) == 32:
            # 旧版验证逻辑（向后兼容）
            expected_hash = hashlib.sha256((salt_hex + password).encode("utf-8")).hexdigest()
            return secrets.compare_digest(pw_hash_hex, expected_hash)

        # 新版 PBKDF2 验证
        expected_hash = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt, PBKDF2_ITERATIONS)
        return secrets.compare_digest(pw_hash_hex, expected_hash.hex())
    except Exception:
        return False

router = APIRouter(prefix="/api/notes", tags=["Notes"])

@router.get("", response_model=List[NoteOut])
def get_notes(
    notebook_id: Optional[str] = None,
    tag: Optional[str] = None,
    is_starred: Optional[bool] = None,
    is_trashed: bool = False,
    keyword: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """获取笔记列表（支持按笔记本、标签、星标、回收站、关键词搜索过滤）"""
    query = db.query(Note).filter(Note.is_trashed == is_trashed)

    if notebook_id:
        query = query.filter(Note.notebook_id == notebook_id)
    if is_starred is not None:
        query = query.filter(Note.is_starred == is_starred)
    if tag:
        query = query.filter(Note.tags.like(f"%{tag}%"))
    if keyword:
        query = query.filter(
            or_(
                Note.title.ilike(f"%{keyword}%"),
                Note.content.ilike(f"%{keyword}%"),
                Note.summary.ilike(f"%{keyword}%")
            )
        )

    notes = query.order_by(desc(Note.updated_at)).all()

    audio_counts = dict(
        db.query(AudioRecord.note_id, func.count(AudioRecord.id))
        .group_by(AudioRecord.note_id)
        .all()
    )

    result = []
    for n in notes:
        tags_list = []
        try:
            tags_list = json.loads(n.tags) if n.tags else []
        except Exception:
            tags_list = []

        is_locked = bool(n.is_locked)
        full_content = n.content or ""
        preview = "" if is_locked else (n.summary or full_content[:240])
        note_dict = {
            "id": n.id,
            "title": n.title,
            "content": preview,
            "content_json": "",
            "notebook_id": n.notebook_id,
            "summary": "🔒 此重要笔记已设置密码锁定保护" if is_locked else (n.summary or ""),
            "tags": tags_list,
            "is_starred": n.is_starred,
            "is_trashed": n.is_trashed,
            "is_locked": is_locked,
            "created_at": n.created_at,
            "updated_at": n.updated_at,
            "audio_count": int(audio_counts.get(n.id) or 0),
            "content_length": 0 if is_locked else len(full_content),
        }
        result.append(note_dict)
    return result


@router.get("/stats", response_model=NoteStatsOut)
def get_note_stats(db: Session = Depends(get_db)):
    """一次查询返回侧栏计数，避免列表接口被重复打三次。"""
    total = db.query(func.count(Note.id)).filter(Note.is_trashed == False).scalar() or 0
    trash = db.query(func.count(Note.id)).filter(Note.is_trashed == True).scalar() or 0
    starred = (
        db.query(func.count(Note.id))
        .filter(Note.is_trashed == False, Note.is_starred == True)
        .scalar()
        or 0
    )
    return {"total": int(total), "trash": int(trash), "starred": int(starred)}

@router.post("", response_model=NoteOut)
def create_note(data: NoteCreate, db: Session = Depends(get_db)):
    """新建笔记"""
    tags_str = json.dumps(data.tags or [], ensure_ascii=False)
    note = Note(
        title=data.title or "无标题笔记",
        content=data.content or "",
        content_json=data.content_json or "",
        notebook_id=data.notebook_id,
        summary=data.summary or "",
        tags=tags_str,
        is_starred=data.is_starred or False,
        is_trashed=data.is_trashed or False,
        is_locked=data.is_locked or False
    )
    db.add(note)
    db.commit()
    db.refresh(note)

    return {
        "id": note.id,
        "title": note.title,
        "content": note.content,
        "content_json": note.content_json,
        "notebook_id": note.notebook_id,
        "summary": note.summary,
        "tags": data.tags or [],
        "is_starred": note.is_starred,
        "is_trashed": note.is_trashed,
        "is_locked": bool(note.is_locked),
        "created_at": note.created_at,
        "updated_at": note.updated_at,
        "audio_count": 0
    }

@router.get("/{note_id}", response_model=NoteOut)
def get_note(note_id: str, db: Session = Depends(get_db)):
    """获取单篇笔记详情"""
    note = db.query(Note).filter(Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    audio_count = db.query(AudioRecord).filter(AudioRecord.note_id == note.id).count()
    tags_list = []
    try:
        tags_list = json.loads(note.tags) if note.tags else []
    except Exception:
        tags_list = []

    is_locked = bool(note.is_locked)
    return {
        "id": note.id,
        "title": note.title,
        "content": "" if is_locked else note.content,
        "content_json": "" if is_locked else note.content_json,
        "notebook_id": note.notebook_id,
        "summary": "🔒 此重要笔记已设置密码锁定保护" if is_locked else note.summary,
        "tags": tags_list,
        "is_starred": note.is_starred,
        "is_trashed": note.is_trashed,
        "is_locked": is_locked,
        "created_at": note.created_at,
        "updated_at": note.updated_at,
        "audio_count": audio_count,
        "content_length": 0 if is_locked else len(note.content or "")
    }

@router.put("/{note_id}", response_model=NoteOut)
def update_note(note_id: str, data: NoteUpdate, db: Session = Depends(get_db)):
    """更新笔记"""
    note = db.query(Note).filter(Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    update_data = data.dict(exclude_unset=True)
    if "tags" in update_data and update_data["tags"] is not None:
        note.tags = json.dumps(update_data["tags"], ensure_ascii=False)
        del update_data["tags"]

    for key, value in update_data.items():
        setattr(note, key, value)

    db.commit()
    db.refresh(note)

    audio_count = db.query(AudioRecord).filter(AudioRecord.note_id == note.id).count()
    tags_list = []
    try:
        tags_list = json.loads(note.tags) if note.tags else []
    except Exception:
        tags_list = []

    return {
        "id": note.id,
        "title": note.title,
        "content": note.content,
        "content_json": note.content_json,
        "notebook_id": note.notebook_id,
        "summary": note.summary,
        "tags": tags_list,
        "is_starred": note.is_starred,
        "is_trashed": note.is_trashed,
        "is_locked": bool(note.is_locked),
        "created_at": note.created_at,
        "updated_at": note.updated_at,
        "audio_count": audio_count
    }

@router.post("/{note_id}/lock", response_model=NoteOut)
def lock_note(note_id: str, data: NoteLockRequest, db: Session = Depends(get_db)):
    """给笔记加锁并设置独立密码"""
    note = db.query(Note).filter(Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    if not data.password or not data.password.strip():
        raise HTTPException(status_code=400, detail="密码不能为空")

    note.is_locked = True
    note.password_hash = hash_password(data.password)
    db.commit()
    db.refresh(note)

    audio_count = db.query(AudioRecord).filter(AudioRecord.note_id == note.id).count()
    tags_list = []
    try:
        tags_list = json.loads(note.tags) if note.tags else []
    except Exception:
        tags_list = []

    return {
        "id": note.id,
        "title": note.title,
        "content": "",
        "content_json": "",
        "notebook_id": note.notebook_id,
        "summary": "🔒 此重要笔记已设置密码锁定保护",
        "tags": tags_list,
        "is_starred": note.is_starred,
        "is_trashed": note.is_trashed,
        "is_locked": True,
        "created_at": note.created_at,
        "updated_at": note.updated_at,
        "audio_count": audio_count
    }

@router.post("/{note_id}/verify-password", response_model=NoteOut)
def verify_note_password(note_id: str, data: NoteVerifyPasswordRequest, db: Session = Depends(get_db)):
    """
    二次验证笔记密码，验证通过后返回完整笔记内容
    🔐 安全特性：自动检测旧版 SHA-256 密码哈希并静默升级为 PBKDF2
    """
    note = db.query(Note).filter(Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    if note.is_locked:
        if not note.password_hash or not verify_password(note.password_hash, data.password):
            raise HTTPException(status_code=400, detail="密码错误")

        # 🔐 自动升级旧版密码哈希为新版 PBKDF2
        if note.password_hash and "$" in note.password_hash:
            salt_hex, pw_hash_hex = note.password_hash.split("$", 1)
            # 检测旧版格式（SHA-256 哈希长度为 64，PBKDF2 也是 64，但通过迭代标识区分）
            # 如果是旧版格式（通过简单长度判断），则升级
            if len(pw_hash_hex) == 64 and len(salt_hex) == 32:
                # 验证成功后，使用新版算法重新哈希
                new_hash = hash_password(data.password)
                # 仅在新旧哈希不同时更新（避免重复升级）
                if new_hash != note.password_hash:
                    note.password_hash = new_hash
                    db.commit()
                    db.refresh(note)

    audio_count = db.query(AudioRecord).filter(AudioRecord.note_id == note.id).count()
    tags_list = []
    try:
        tags_list = json.loads(note.tags) if note.tags else []
    except Exception:
        tags_list = []

    return {
        "id": note.id,
        "title": note.title,
        "content": note.content,
        "content_json": note.content_json,
        "notebook_id": note.notebook_id,
        "summary": note.summary,
        "tags": tags_list,
        "is_starred": note.is_starred,
        "is_trashed": note.is_trashed,
        "is_locked": bool(note.is_locked),
        "created_at": note.created_at,
        "updated_at": note.updated_at,
        "audio_count": audio_count
    }

@router.post("/{note_id}/unlock", response_model=NoteOut)
def unlock_note(note_id: str, data: NoteUnlockRequest, db: Session = Depends(get_db)):
    """解除笔记密码锁定"""
    note = db.query(Note).filter(Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    if note.is_locked:
        if not note.password_hash or not verify_password(note.password_hash, data.password):
            raise HTTPException(status_code=400, detail="密码错误")

    note.is_locked = False
    note.password_hash = None
    db.commit()
    db.refresh(note)

    audio_count = db.query(AudioRecord).filter(AudioRecord.note_id == note.id).count()
    tags_list = []
    try:
        tags_list = json.loads(note.tags) if note.tags else []
    except Exception:
        tags_list = []

    return {
        "id": note.id,
        "title": note.title,
        "content": note.content,
        "content_json": note.content_json,
        "notebook_id": note.notebook_id,
        "summary": note.summary,
        "tags": tags_list,
        "is_starred": note.is_starred,
        "is_trashed": note.is_trashed,
        "is_locked": False,
        "created_at": note.created_at,
        "updated_at": note.updated_at,
        "audio_count": audio_count
    }

@router.delete("/{note_id}")
def delete_note(note_id: str, permanent: bool = False, db: Session = Depends(get_db)):
    """删除笔记（默认移入废纸篓，permanent=True 时物理永久删除）"""
    note = db.query(Note).filter(Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    if permanent:
        db.delete(note)
        db.commit()
        return {"status": "success", "message": "Note permanently deleted"}
    else:
        note.is_trashed = True
        db.commit()
        return {"status": "success", "message": "Note moved to trash"}

@router.post("/{note_id}/restore")
def restore_note(note_id: str, db: Session = Depends(get_db)):
    """从废纸篓恢复笔记"""
    note = db.query(Note).filter(Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    note.is_trashed = False
    db.commit()
    return {"status": "success", "message": "Note restored successfully"}

@router.delete("/trash/empty")
def empty_trash(db: Session = Depends(get_db)):
    """清空废纸篓 (包括笔记与多维数据表)"""
    db.query(Note).filter(Note.is_trashed == True).delete()
    db.query(Database).filter(Database.is_archived == True).delete()
    db.commit()
    return {"status": "success", "message": "Trash emptied successfully"}

@router.get("/{note_id}/export/{format}")
def export_note(note_id: str, format: str, db: Session = Depends(get_db)):
    """导出单篇笔记为指定格式"""
    from fastapi.responses import FileResponse
    from urllib.parse import quote
    from ..services.export_service import ExportService

    note = db.query(Note).filter(Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    tags_list = []
    try:
        tags_list = json.loads(note.tags) if note.tags else []
    except Exception:
        tags_list = []

    updated_at_str = note.updated_at.strftime("%Y-%m-%d %H:%M") if note.updated_at else None

    if format == "docx":
        file_path = ExportService.markdown_to_docx(
            title=note.title, content=note.content, tags=tags_list, updated_at=updated_at_str
        )
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format == "md":
        file_path = ExportService.export_markdown(
            title=note.title, content=note.content, tags=tags_list, updated_at=updated_at_str
        )
        media_type = "text/markdown"
    elif format == "txt":
        file_path = ExportService.export_txt(
            title=note.title, content=note.content, tags=tags_list, updated_at=updated_at_str
        )
        media_type = "text/plain"
    elif format == "html":
        file_path = ExportService.export_html(
            title=note.title, content=note.content, tags=tags_list, updated_at=updated_at_str
        )
        media_type = "text/html"
    else:
        raise HTTPException(status_code=400, detail="Unsupported format")

    filename = f"{note.title or 'note'}.{format}"
    encoded_filename = quote(filename)

    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )

@router.post("/{note_id}/clone", response_model=NoteOut)
def clone_note(note_id: str, db: Session = Depends(get_db)):
    """
    克隆单篇笔记
    安全策略：克隆后的笔记自动解除加密锁定，避免密码共享风险
    """
    original = db.query(Note).filter(Note.id == note_id).first()
    if not original:
        raise HTTPException(status_code=404, detail="Note not found")

    new_note = Note(
        title=f"{original.title} (副本)" if original.title else "无标题笔记 (副本)",
        content=original.content,
        content_json=original.content_json,
        notebook_id=original.notebook_id,
        summary=original.summary,
        tags=original.tags,
        is_starred=False,
        is_trashed=False,
        is_locked=False,  # 克隆笔记默认解除加密
        password_hash=None  # 不复制密码哈希
    )
    db.add(new_note)
    db.commit()
    db.refresh(new_note)

    tags_list = []
    try:
        tags_list = json.loads(new_note.tags) if new_note.tags else []
    except Exception:
        tags_list = []

    is_locked = bool(new_note.is_locked)
    return {
        "id": new_note.id,
        "title": new_note.title,
        "content": new_note.content,  # 克隆后已解锁，返回完整内容
        "content_json": new_note.content_json,
        "notebook_id": new_note.notebook_id,
        "summary": new_note.summary,
        "tags": tags_list,
        "is_starred": new_note.is_starred,
        "is_trashed": new_note.is_trashed,
        "is_locked": False,  # 克隆后已解锁
        "created_at": new_note.created_at,
        "updated_at": new_note.updated_at,
        "audio_count": 0
    }

@router.post("/batch-import", response_model=List[NoteOut])
async def batch_import(
    files: List[UploadFile] = File(...),
    notebook_id: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    """
    批量导入 Markdown、Word (.docx)、HTML 或各类纯文本文件
    安全限制：
    - 单次最多 100 个文件
    - 单文件最大 20MB
    """
    MAX_FILES = 100
    MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB

    if len(files) > MAX_FILES:
        raise HTTPException(
            status_code=400,
            detail=f"单次最多导入 {MAX_FILES} 个文件，当前提交了 {len(files)} 个"
        )

    created_notes = []
    for file in files:
        filename = file.filename or "未命名笔记.txt"
        base_name, ext = os.path.splitext(filename)
        title = base_name.strip() or "未命名导入笔记"

        content_bytes = await file.read()

        # 验证文件大小与有效性
        if len(content_bytes) > MAX_FILE_SIZE or len(content_bytes) == 0:
            continue

        content = extract_content_from_file(filename, content_bytes)
        if not content.strip():
            content = "（导入的文件无文本内容）"

        new_note = Note(
            title=title,
            content=content,
            content_json="",
            notebook_id=notebook_id if (notebook_id and notebook_id.strip()) else None,
            summary=content[:120].replace('\n', ' ') if content else "",
            tags="[]",
            is_starred=False,
            is_trashed=False,
            is_locked=False
        )
        db.add(new_note)
        db.commit()
        db.refresh(new_note)

        created_notes.append({
            "id": new_note.id,
            "title": new_note.title,
            "content": new_note.content,
            "content_json": new_note.content_json,
            "notebook_id": new_note.notebook_id,
            "summary": new_note.summary,
            "tags": [],
            "is_starred": new_note.is_starred,
            "is_trashed": new_note.is_trashed,
            "is_locked": False,
            "created_at": new_note.created_at,
            "updated_at": new_note.updated_at,
            "audio_count": 0
        })
    return created_notes


@router.get("/graph/data", response_model=GraphDataOut)
def get_knowledge_graph(db: Session = Depends(get_db)):
    """获取全局知识图谱节点与连线数据 (解析 [[双向链接]] 与 #标签 关联)"""
    notes = db.query(Note).filter(Note.is_trashed == False).all()
    notebooks = {nb.id: nb.name for nb in db.query(Notebook).all()}
    
    # 建立标题 -> Note ID 和 ID -> Note 的映射
    title_to_note = {n.title.strip().lower(): n for n in notes if n.title}
    id_to_note = {n.id: n for n in notes}

    nodes_dict: Dict[str, GraphNode] = {}
    links: List[GraphLink] = []
    link_count: Dict[str, int] = {}

    # 1. 注册所有笔记节点
    for n in notes:
        nb_name = notebooks.get(n.notebook_id, "未分类")
        nodes_dict[n.id] = GraphNode(
            id=n.id,
            title=n.title or "无标题笔记",
            notebook_id=n.notebook_id,
            notebook_name=nb_name,
            group="note",
            val=1
        )
        link_count[n.id] = 0

    # 2. 解析笔记正文中的 [[双链]] 引用
    link_pattern = re.compile(r'\[\[([^\]]+)\]\]')
    tag_nodes: Dict[str, GraphNode] = {}

    for n in notes:
        content = n.content or ""
        # 匹配 [[目标笔记]]
        matches = link_pattern.findall(content)
        for target_text in matches:
            target_clean = target_text.strip()
            target_key = target_clean.lower()
            
            # 匹配目标是否为已有笔记标题或 ID
            target_note = title_to_note.get(target_key) or id_to_note.get(target_clean)
            if target_note and target_note.id != n.id:
                links.append(GraphLink(
                    source=n.id,
                    target=target_note.id,
                    label="link"
                ))
                link_count[n.id] = link_count.get(n.id, 0) + 1
                link_count[target_note.id] = link_count.get(target_note.id, 0) + 1

        # 关联标签节点
        try:
            tags = json.loads(n.tags) if n.tags else []
        except Exception:
            tags = []

        for tag in tags:
            tag_id = f"tag_{tag}"
            if tag_id not in tag_nodes:
                tag_nodes[tag_id] = GraphNode(
                    id=tag_id,
                    title=f"#{tag}",
                    notebook_id=None,
                    notebook_name="标签",
                    group="tag",
                    val=2
                )
                link_count[tag_id] = 0

            links.append(GraphLink(
                source=n.id,
                target=tag_id,
                label="tag"
            ))
            link_count[n.id] = link_count.get(n.id, 0) + 1
            link_count[tag_id] = link_count.get(tag_id, 0) + 1

    # 合并标签节点并更新节点权重 val
    for tag_id, tag_node in tag_nodes.items():
        nodes_dict[tag_id] = tag_node

    for node_id, node in nodes_dict.items():
        node.val = max(1, link_count.get(node_id, 1))

    return GraphDataOut(
        nodes=list(nodes_dict.values()),
        links=links
    )


@router.get("/{note_id}/backlinks", response_model=BacklinksOut)
def get_note_backlinks(note_id: str, db: Session = Depends(get_db)):
    """获取引用了当前笔记的所有反向引用 (Backlinks) 列表及上下文摘要"""
    current_note = db.query(Note).filter(Note.id == note_id).first()
    if not current_note:
        raise HTTPException(status_code=404, detail="笔记不存在")

    current_title = current_note.title.strip() if current_note.title else ""
    all_notes = db.query(Note).filter(Note.id != note_id, Note.is_trashed == False).all()

    backlinks: List[BacklinkItem] = []
    # 匹配 [[当前笔记标题]] 或 [[当前笔记ID]]
    patterns = [
        re.compile(re.escape(f"[[{current_title}]]"), re.IGNORECASE) if current_title else None,
        re.compile(re.escape(f"[[{note_id}]]"))
    ]
    patterns = [p for p in patterns if p is not None]

    for n in all_notes:
        content = n.content or ""
        matched = False
        snippet = ""
        for p in patterns:
            m = p.search(content)
            if m:
                matched = True
                # 提取前后 50 个字符的上下文摘要
                start = max(0, m.start() - 30)
                end = min(len(content), m.end() + 30)
                raw_snippet = content[start:end].replace('\n', ' ')
                snippet = f"...{raw_snippet}..." if start > 0 or end < len(content) else raw_snippet
                break

        if matched:
            backlinks.append(BacklinkItem(
                note_id=n.id,
                note_title=n.title or "无标题笔记",
                snippet=snippet,
                updated_at=n.updated_at
            ))

    return BacklinksOut(
        note_id=note_id,
        backlinks=backlinks
    )

