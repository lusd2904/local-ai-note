import os
import re
from pathlib import Path
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """设置 Word 表格单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

class ExportService:

    @staticmethod
    def markdown_to_docx(title: str, content: str, tags: list = None, updated_at: str = None) -> Path:
        """将笔记内容转换为排版精美的 Word 文档 (.docx) 并保存到 exports 目录"""
        from ..config import EXPORTS_DIR
        export_dir = EXPORTS_DIR
        export_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # 设置文档页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # 1. 标题 (Heading 1)
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title or "无标题笔记")
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 41, 55) # 深色
        title_p.paragraph_format.space_after = Pt(8)

        # 2. 元信息（更新时间与标签）
        meta_p = doc.add_paragraph()
        meta_text = f"更新时间: {updated_at or datetime.now().strftime('%Y-%m-%d %H:%M')}"
        if tags and len(tags) > 0:
            meta_text += f"   |   标签: {', '.join(tags)}"
        meta_run = meta_p.add_run(meta_text)
        meta_run.font.size = Pt(9.5)
        meta_run.font.color.rgb = RGBColor(107, 114, 128) # 灰色
        meta_p.paragraph_format.space_after = Pt(16)

        # 分割线
        divider_p = doc.add_paragraph()
        divider_run = divider_p.add_run("―" * 40)
        divider_run.font.color.rgb = RGBColor(229, 231, 235)
        divider_p.paragraph_format.space_after = Pt(14)

        # 3. 逐行解析 Markdown 正文
        lines = (content or "").split("\n")
        in_code_block = False
        code_block_lines = []

        for line in lines:
            # 代码块
            if line.startswith("```"):
                if in_code_block:
                    # 结束代码块
                    code_text = "\n".join(code_block_lines)
                    table = doc.add_table(rows=1, cols=1)
                    table.autofit = False
                    table.columns[0].width = Inches(6.5)
                    cell = table.cell(0, 0)
                    set_cell_background(cell, "F3F4F6")
                    cp = cell.paragraphs[0]
                    crun = cp.add_run(code_text)
                    crun.font.size = Pt(9.5)
                    crun.font.name = "Courier New"
                    crun.font.color.rgb = RGBColor(55, 65, 81)
                    doc.add_paragraph().paragraph_format.space_after = Pt(4)
                    code_block_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_block_lines = []
                continue

            if in_code_block:
                code_block_lines.append(line)
                continue

            # 空行
            if not line.strip():
                continue

            # 标题处理 (#, ##, ###)
            if line.startswith("# "):
                h = doc.add_paragraph()
                hrun = h.add_run(line[2:].strip())
                hrun.font.size = Pt(16)
                hrun.font.bold = True
                hrun.font.color.rgb = RGBColor(37, 99, 235) # 蓝色
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
            elif line.startswith("## "):
                h = doc.add_paragraph()
                hrun = h.add_run(line[3:].strip())
                hrun.font.size = Pt(14)
                hrun.font.bold = True
                hrun.font.color.rgb = RGBColor(31, 41, 55)
                h.paragraph_format.space_before = Pt(10)
                h.paragraph_format.space_after = Pt(4)
            elif line.startswith("### "):
                h = doc.add_paragraph()
                hrun = h.add_run(line[4:].strip())
                hrun.font.size = Pt(12)
                hrun.font.bold = True
                hrun.font.color.rgb = RGBColor(55, 65, 81)
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after = Pt(4)
            elif line.startswith("> "):
                # 引用块
                qp = doc.add_paragraph()
                qp.paragraph_format.left_indent = Inches(0.3)
                qrun = qp.add_run(line[2:].strip())
                qrun.font.italic = True
                qrun.font.size = Pt(10.5)
                qrun.font.color.rgb = RGBColor(75, 85, 99)
                qp.paragraph_format.space_after = Pt(6)
            elif line.startswith("- ") or line.startswith("* "):
                # 无序列表
                lp = doc.add_paragraph(style='List Bullet')
                text = line[2:].strip()
                # 简单解析 **加粗**
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = lp.add_run(part[2:-2])
                        r.font.bold = True
                    else:
                        lp.add_run(part)
                lp.paragraph_format.space_after = Pt(3)
            elif re.match(r'^\d+\.\s', line):
                # 有序列表
                lp = doc.add_paragraph(style='List Number')
                text = re.sub(r'^\d+\.\s', '', line).strip()
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = lp.add_run(part[2:-2])
                        r.font.bold = True
                    else:
                        lp.add_run(part)
                lp.paragraph_format.space_after = Pt(3)
            else:
                # 普通段落
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2])
                        r.font.bold = True
                    else:
                        p.add_run(part)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.25

        # 生成安全文件名并保存
        safe_title = re.sub(r'[\/\\:\*\?"<>\|]', '_', title or "note")[:50]
        file_path = export_dir / f"{safe_title}.docx"
        doc.save(str(file_path))
        return file_path

    @staticmethod
    def export_markdown(title: str, content: str, tags: list = None, updated_at: str = None) -> Path:
        from ..config import EXPORTS_DIR
        export_dir = EXPORTS_DIR
        export_dir.mkdir(parents=True, exist_ok=True)
        safe_title = re.sub(r'[\/\\:\*\?"<>\|]', '_', title or "note")[:50]
        file_path = export_dir / f"{safe_title}.md"
        
        tags_str = f"[{', '.join(tags)}]" if tags else "[]"
        date_str = updated_at or datetime.now().strftime('%Y-%m-%d %H:%M')
        
        frontmatter = f"---\ntitle: {title}\ndate: {date_str}\ntags: {tags_str}\n---\n\n"
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(frontmatter + (content or ""))
            
        return file_path

    @staticmethod
    def export_txt(title: str, content: str, tags: list = None, updated_at: str = None) -> Path:
        from ..config import EXPORTS_DIR
        export_dir = EXPORTS_DIR
        export_dir.mkdir(parents=True, exist_ok=True)
        safe_title = re.sub(r'[\/\\:\*\?"<>\|]', '_', title or "note")[:50]
        file_path = export_dir / f"{safe_title}.txt"
        
        header = f"标题: {title}\n更新时间: {updated_at or datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        if tags:
            header += f"标签: {', '.join(tags)}\n"
        header += "-" * 40 + "\n\n"
        
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(header + (content or ""))
            
        return file_path

    @staticmethod
    def export_html(title: str, content: str, tags: list = None, updated_at: str = None) -> Path:
        from ..config import EXPORTS_DIR
        export_dir = EXPORTS_DIR
        export_dir.mkdir(parents=True, exist_ok=True)
        safe_title = re.sub(r'[\/\\:\*\?"<>\|]', '_', title or "note")[:50]
        file_path = export_dir / f"{safe_title}.html"

        try:
            import markdown
            html_content = markdown.markdown(content or "", extensions=['extra', 'codehilite'])
        except Exception:
            import html as html_lib
            escaped = html_lib.escape(content or "")
            html_content = f"<pre style='white-space: pre-wrap;'>{escaped}</pre>"
        
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; }}
        .prose pre {{ background-color: #f3f4f6; padding: 1rem; border-radius: 0.375rem; overflow-x: auto; }}
        .prose code {{ font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", "Courier New", monospace; }}
    </style>
</head>
<body class="bg-gray-50 py-8">
    <div class="max-w-4xl mx-auto bg-white p-8 rounded-lg shadow-sm">
        <h1 class="text-3xl font-bold text-gray-900 mb-4">{title}</h1>
        <div class="text-sm text-gray-500 mb-8 pb-4 border-b">
            <span>更新时间: {updated_at or datetime.now().strftime('%Y-%m-%d %H:%M')}</span>
"""
        if tags:
            html += f"""            <span class="ml-4">标签: {', '.join(tags)}</span>\n"""
        
        html += f"""        </div>
        <article class="prose max-w-none">
            {html_content}
        </article>
    </div>
</body>
</html>"""
        
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(html)
            
        return file_path
